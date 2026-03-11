import re
from dataclasses import dataclass
from io import BytesIO
from typing import Dict, List, Tuple

import pandas as pd
import streamlit as st
from docx import Document
from pypdf import PdfReader


st.set_page_config(page_title="AAEL VALUE Rubric Generator", page_icon="📘", layout="wide")


@dataclass
class ProblemInfo:
    number: int
    title: str
    body: str
    problem_type: str
    suggested_points: int


VALUE_RUBRICS: Dict[str, List[Dict[str, str]]] = {
    "Quantitative Literacy": [
        {"criterion": "Interpretation"},
        {"criterion": "Representation"},
        {"criterion": "Calculation"},
        {"criterion": "Application/Analysis"},
        {"criterion": "Assumptions"},
        {"criterion": "Communication"},
    ],
    "Inquiry and Analysis": [
        {"criterion": "Topic Selection"},
        {"criterion": "Existing Knowledge, Research, and/or Views"},
        {"criterion": "Design Process"},
        {"criterion": "Analysis"},
        {"criterion": "Conclusions"},
        {"criterion": "Limitations and Implications"},
    ],
    "Problem Solving": [
        {"criterion": "Define Problem"},
        {"criterion": "Identify Strategies"},
        {"criterion": "Propose Solutions/Hypotheses"},
        {"criterion": "Evaluate Potential Solutions"},
        {"criterion": "Implement Solution"},
        {"criterion": "Evaluate Outcomes"},
    ],
    "Written Communication": [
        {"criterion": "Context of and Purpose for Writing"},
        {"criterion": "Content Development"},
        {"criterion": "Genre and Disciplinary Conventions"},
        {"criterion": "Sources and Evidence"},
        {"criterion": "Control of Syntax and Mechanics"},
    ],
}

DEFAULT_WEIGHTS = {
    "Quantitative Literacy": 35,
    "Inquiry and Analysis": 25,
    "Problem Solving": 20,
    "Written Communication": 20,
}


def extract_pdf_text(uploaded_file) -> str:
    reader = PdfReader(uploaded_file)
    return "\n".join((page.extract_text() or "") for page in reader.pages)



def clean_text(text: str) -> str:
    text = text.replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    return text.strip()



def infer_assignment_name(text: str) -> str:
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    if not lines:
        return ""

    bad_starts = (
        "problem ",
        "questions",
        "requirements",
        "submission",
        "tasks",
        "objective",
        "instructions",
    )
    for line in lines[:10]:
        lowered = line.lower()
        if len(line) <= 90 and not lowered.startswith(bad_starts):
            return line
    return ""



def infer_total_points(text: str, num_problems: int) -> int | None:
    total_match = re.search(r"total\s+points\s*[:\-]?\s*(\d+)", text, flags=re.I)
    if total_match:
        return int(total_match.group(1))

    worth_match = re.search(r"worth\s+(\d+)\s+points", text, flags=re.I)
    if worth_match:
        return int(worth_match.group(1))

    problem_points = [int(x) for x in re.findall(r"\((\d+)\s*points?\)", text, flags=re.I)]
    if problem_points and len(problem_points) >= max(1, num_problems):
        return sum(problem_points[:num_problems])
    return None



def split_into_problems(text: str) -> List[Tuple[str, str]]:
    pattern = re.compile(r"(?im)^\s*(problem\s+\d+[^\n]*)")
    matches = list(pattern.finditer(text))
    if not matches:
        return [("Full Assignment", text)]

    sections: List[Tuple[str, str]] = []
    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        sections.append((match.group(1).strip(), text[start:end].strip()))
    return sections



def infer_problem_type(body: str) -> str:
    lower = body.lower()
    r_markers = [" in r", "using r", " lm(", "ggplot", "ggplot2", "r code"]
    python_markers = ["python", "pandas", "numpy", "statsmodels", "scikit", "matplotlib", "python code"]
    coding_markers = ["code", "script", "program", "plot", "regression line", "fit a model"]
    written_markers = ["explain", "interpret", "in your own words", "show your calculation", "plain language", "conceptual"]

    has_r = any(marker in lower for marker in r_markers)
    has_python = any(marker in lower for marker in python_markers)
    has_coding = has_r or has_python or any(marker in lower for marker in coding_markers)
    has_written = any(marker in lower for marker in written_markers)

    if has_r and not has_python:
        return "R coding"
    if has_python and not has_r:
        return "Python coding"
    if has_r and has_python:
        return "Mixed"
    if has_coding and not has_written:
        return "Coding"
    if has_written and not has_coding:
        return "Written analysis"
    if has_written and has_coding:
        return "Mixed"
    return "Written analysis"



def suggest_points(num_problems: int, total_points: int) -> List[int]:
    if num_problems <= 0:
        return []
    base = total_points // num_problems
    remainder = total_points % num_problems
    return [base + (1 if i < remainder else 0) for i in range(num_problems)]



def parse_problems(text: str, total_points: int) -> List[ProblemInfo]:
    sections = split_into_problems(text)
    suggested = suggest_points(len(sections), total_points)
    problems: List[ProblemInfo] = []

    for idx, (heading, body) in enumerate(sections, start=1):
        number_match = re.search(r"problem\s+(\d+)", heading, flags=re.I)
        number = int(number_match.group(1)) if number_match else idx
        problems.append(
            ProblemInfo(
                number=number,
                title=heading,
                body=body,
                problem_type=infer_problem_type(body),
                suggested_points=suggested[idx - 1],
            )
        )
    return problems



def infer_default_value_set(text: str) -> List[str]:
    lower = text.lower()
    selected = ["Quantitative Literacy", "Inquiry and Analysis", "Written Communication"]
    if any(k in lower for k in ["python", "using python", "using r", "r code", "script", "program"]):
        selected.insert(2, "Problem Solving")
    return selected



def normalize_weights(weight_map: Dict[str, int], total_points: int) -> Dict[str, int]:
    selected = {k: v for k, v in weight_map.items() if v > 0}
    if not selected:
        return {}

    total_weight = sum(selected.values())
    out: Dict[str, int] = {}
    running = 0
    keys = list(selected.keys())
    for i, key in enumerate(keys):
        if i < len(keys) - 1:
            pts = round(total_points * selected[key] / total_weight)
            running += pts
        else:
            pts = total_points - running
        out[key] = pts
    return out



def extract_topic_keywords(text: str) -> List[str]:
    candidates = [
        "linear regression",
        "logistic regression",
        "multiple regression",
        "prediction",
        "model fit",
        "coefficients",
        "python",
        "r",
        "written interpretation",
        "visualization",
        "scatter plot",
        "regression output",
        "business context",
    ]
    lower = text.lower()
    found = [c for c in candidates if c in lower]
    return found[:5]



def descriptor_focus(rubric_name: str, criterion: str, assignment_label: str, keywords: List[str]) -> str:
    topic = keywords[0] if keywords else "the assignment"
    coding_present = any(k in keywords for k in ["python", "r"])

    custom_map = {
        ("Quantitative Literacy", "Interpretation"): f"Explain {topic} equations, output, and graphs accurately in the context of {assignment_label}.",
        ("Quantitative Literacy", "Representation"): f"Translate results from {topic} into equations, tables, plots, or code-based output.",
        ("Quantitative Literacy", "Calculation"): f"Carry out calculations and predictions correctly for {assignment_label}.",
        ("Quantitative Literacy", "Application/Analysis"): f"Use quantitative evidence from {topic} to support reasonable conclusions.",
        ("Quantitative Literacy", "Assumptions"): f"Identify the assumptions or limits that affect confidence in the {topic} results.",
        ("Quantitative Literacy", "Communication"): f"Present quantitative evidence clearly through writing, tables, plots, and outputs for {assignment_label}.",
        ("Inquiry and Analysis", "Topic Selection"): f"Frame the core analytical problem clearly and keep the work focused on {assignment_label}.",
        ("Inquiry and Analysis", "Existing Knowledge, Research, and/or Views"): f"Use relevant concepts from the course materials to support the {topic} analysis.",
        ("Inquiry and Analysis", "Design Process"): f"Choose an appropriate workflow, including written reasoning and coding steps where required.",
        ("Inquiry and Analysis", "Analysis"): f"Organize evidence to reveal meaningful relationships, fit, and patterns in the {topic} task.",
        ("Inquiry and Analysis", "Conclusions"): f"State conclusions that follow logically from the results in {assignment_label}.",
        ("Inquiry and Analysis", "Limitations and Implications"): f"Discuss limits of the model and what the findings do or do not imply.",
        ("Problem Solving", "Define Problem"): f"Identify the problem clearly and connect it to the goals of {assignment_label}.",
        ("Problem Solving", "Identify Strategies"): f"Select appropriate written and computational strategies for solving the task.",
        ("Problem Solving", "Propose Solutions/Hypotheses"): f"Develop a reasonable modeling or solution approach for the problem.",
        ("Problem Solving", "Evaluate Potential Solutions"): f"Justify why the chosen method or interpretation is appropriate for the task.",
        ("Problem Solving", "Implement Solution"): f"Carry out the solution effectively, including correct code, output, and written work where needed.",
        ("Problem Solving", "Evaluate Outcomes"): f"Review the results thoughtfully and indicate whether further work or refinement is needed.",
        ("Written Communication", "Context of and Purpose for Writing"): f"Respond clearly to the assignment directions, audience, and purpose.",
        ("Written Communication", "Content Development"): f"Develop explanations and interpretations fully using relevant details from {assignment_label}.",
        ("Written Communication", "Genre and Disciplinary Conventions"): f"Follow conventions of business analytics, statistics, and course formatting expectations.",
        ("Written Communication", "Sources and Evidence"): f"Use calculations, outputs, plots, and chapter-based evidence appropriately.",
        ("Written Communication", "Control of Syntax and Mechanics"): "Write clearly, professionally, and with minimal distracting errors.",
    }
    focus = custom_map.get((rubric_name, criterion), f"Demonstrate strong performance on {criterion.lower()} for {assignment_label}.")
    if coding_present and rubric_name == "Written Communication" and criterion == "Genre and Disciplinary Conventions":
        focus += " Include readable code formatting and comments when code is required."
    return focus



def build_level_descriptors(rubric_name: str, criterion: str, focus: str, problem_types: List[str]) -> Dict[str, str]:
    coding_present = any(pt in {"Python coding", "R coding", "Coding", "Mixed"} for pt in problem_types)
    written_present = any(pt in {"Written analysis", "Mixed"} for pt in problem_types)
    code_phrase = " Code is organized, readable, and appropriately commented." if coding_present else ""
    writing_phrase = " Written explanations are clear and accurate." if written_present else ""

    base4 = f"Consistently exceeds expectations for this criterion. {focus}{code_phrase}{writing_phrase}".strip()
    base3 = f"Meets expectations for this criterion with only minor weaknesses. {focus}".strip()
    base2 = f"Partially meets expectations but shows noticeable gaps in accuracy, completeness, or clarity related to {criterion.lower()}.".strip()
    base1 = f"Does not yet meet expectations. Work is incomplete, inaccurate, or unclear for {criterion.lower()}.".strip()

    if rubric_name == "Quantitative Literacy" and criterion == "Calculation":
        base4 = "Calculations, predictions, and numeric work are accurate, complete, and clearly shown where appropriate."
        base3 = "Calculations and predictions are mostly accurate with only minor omissions or small errors."
        base2 = "Calculations are partly correct but include meaningful errors, omissions, or weak support."
        base1 = "Calculations are largely incorrect, missing, or unsupported."
    elif rubric_name == "Written Communication" and criterion == "Control of Syntax and Mechanics":
        base4 = "Writing is clear, professional, and virtually free of distracting errors."
        base3 = "Writing is clear overall and contains few distracting errors."
        base2 = "Writing communicates meaning but includes recurring wording, grammar, or mechanics issues."
        base1 = "Errors in wording, grammar, or mechanics interfere with clarity and professionalism."
    elif rubric_name == "Problem Solving" and criterion == "Implement Solution" and coding_present:
        base4 = "Implements the solution correctly and efficiently, with clean code, correct output, and strong alignment to the task."
        base3 = "Implements the solution correctly overall, with minor issues in code quality, output, or completeness."
        base2 = "Implements part of the solution but with notable gaps in correctness, output, or readability."
        base1 = "Implementation is incomplete or does not solve the stated problem."

    return {"Level 4": base4, "Level 3": base3, "Level 2": base2, "Level 1": base1}



def build_value_rubric(selected_rubrics: List[str], total_points: int, rubric_weights: Dict[str, int], assignment_label: str, text: str, problem_types: List[str]) -> pd.DataFrame:
    normalized = normalize_weights(rubric_weights, total_points)
    keywords = extract_topic_keywords(text)
    rows = []

    for rubric_name in selected_rubrics:
        criteria = VALUE_RUBRICS[rubric_name]
        rubric_points = normalized.get(rubric_name, 0)
        per_criterion = rubric_points // len(criteria)
        remainder = rubric_points % len(criteria)

        for i, item in enumerate(criteria):
            criterion_points = per_criterion + (1 if i < remainder else 0)
            focus = descriptor_focus(rubric_name, item["criterion"], assignment_label, keywords)
            levels = build_level_descriptors(rubric_name, item["criterion"], focus, problem_types)
            rows.append(
                {
                    "VALUE Rubric": rubric_name,
                    "Criterion": item["criterion"],
                    "Focus": focus,
                    "Points": criterion_points,
                    **levels,
                }
            )
    return pd.DataFrame(rows)



def build_problem_map(problems: List[ProblemInfo], selected_rubrics: List[str]) -> pd.DataFrame:
    rows = []
    for p in problems:
        if p.problem_type in {"Python coding", "R coding", "Coding", "Mixed"}:
            recommended = [r for r in selected_rubrics if r in ["Quantitative Literacy", "Inquiry and Analysis", "Problem Solving", "Written Communication"]]
        else:
            recommended = [r for r in selected_rubrics if r in ["Quantitative Literacy", "Inquiry and Analysis", "Written Communication"]]
        rows.append(
            {
                "Problem": f"Problem {p.number}",
                "Title": p.title,
                "Type": p.problem_type,
                "Suggested Points": p.suggested_points,
                "Recommended VALUE Rubrics": ", ".join(recommended),
            }
        )
    return pd.DataFrame(rows)



def build_grid_view(value_df: pd.DataFrame) -> pd.DataFrame:
    return value_df[["VALUE Rubric", "Criterion", "Points", "Level 4", "Level 3", "Level 2", "Level 1"]].copy()



def generate_markdown_summary(assignment_name: str, total_points: int, problems: List[ProblemInfo], value_df: pd.DataFrame) -> str:
    lines = [f"# AAC&U VALUE-Aligned Rubric for {assignment_name}", "", f"**Total Points:** {total_points}", "", "## Detected Problem Structure", ""]
    for p in problems:
        lines.append(f"- **Problem {p.number}:** {p.problem_type} ({p.suggested_points} suggested points)")
    lines.extend(["", "## VALUE Criteria", ""])

    current = None
    for _, row in value_df.iterrows():
        if row["VALUE Rubric"] != current:
            current = row["VALUE Rubric"]
            lines.append(f"### {current}")
        lines.append(f"**{row['Criterion']} ({row['Points']} pts)**")
        lines.append(f"Focus: {row['Focus']}")
        lines.append(f"- **4:** {row['Level 4']}")
        lines.append(f"- **3:** {row['Level 3']}")
        lines.append(f"- **2:** {row['Level 2']}")
        lines.append(f"- **1:** {row['Level 1']}")
        lines.append("")
    return "\n".join(lines)



def to_csv_bytes(df: pd.DataFrame) -> bytes:
    return df.to_csv(index=False).encode("utf-8")



def to_docx_bytes(assignment_name: str, total_points: int, problem_map_df: pd.DataFrame, value_df: pd.DataFrame) -> bytes:
    doc = Document()
    doc.add_heading(f"AAC&U VALUE-Aligned Rubric: {assignment_name}", level=1)
    doc.add_paragraph(f"Total Points: {total_points}")

    doc.add_heading("Detected Problems", level=2)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Problem"
    hdr[1].text = "Title"
    hdr[2].text = "Type"
    hdr[3].text = "Suggested Points"
    for _, row in problem_map_df.iterrows():
        cells = table.add_row().cells
        cells[0].text = str(row["Problem"])
        cells[1].text = str(row["Title"])
        cells[2].text = str(row["Type"])
        cells[3].text = str(row["Suggested Points"])

    current = None
    for _, row in value_df.iterrows():
        if row["VALUE Rubric"] != current:
            current = row["VALUE Rubric"]
            doc.add_heading(current, level=2)
        doc.add_paragraph(f"{row['Criterion']} ({row['Points']} pts)", style="List Bullet")
        doc.add_paragraph(f"Focus: {row['Focus']}")
        doc.add_paragraph(f"4: {row['Level 4']}")
        doc.add_paragraph(f"3: {row['Level 3']}")
        doc.add_paragraph(f"2: {row['Level 2']}")
        doc.add_paragraph(f"1: {row['Level 1']}")

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()



def init_state() -> None:
    defaults = {
        "source_text": "",
        "detected_assignment_name": "",
        "detected_total_points": None,
        "source_loaded": False,
        "assignment_name": "",
        "total_points": 100,
        "confirmed_metadata": False,
    }
    for key, value in defaults.items():
        if key not in st.session_state:
            st.session_state[key] = value


init_state()

st.title("AAEL VALUE Rubric Generator")
st.caption("Upload an assignment PDF, confirm the metadata, and generate a tailored AAC&U VALUE-aligned rubric.")

with st.sidebar:
    st.header("Settings")
    if st.session_state.confirmed_metadata:
        st.text_input("Assignment name", value=st.session_state.assignment_name, disabled=True)
        st.number_input("Total points", min_value=10, max_value=500, value=int(st.session_state.total_points), disabled=True)
    else:
        st.text_input("Assignment name", value="", placeholder="Will appear after upload or paste", disabled=True)
        st.number_input("Total points", min_value=10, max_value=500, value=100, disabled=True)
    show_raw_text = st.checkbox("Show extracted PDF text", value=False)

st.markdown("### 1 Upload Assignment")
uploaded_file = st.file_uploader("Upload assignment PDF", type=["pdf"])
manual_text = st.text_area("Or paste assignment text here", height=220, placeholder="Paste the full assignment text here if you do not want to upload a PDF.")

if uploaded_file is not None:
    try:
        text = clean_text(extract_pdf_text(uploaded_file))
        st.session_state.source_text = text
        st.session_state.source_loaded = True
        sections = split_into_problems(text)
        inferred_name = infer_assignment_name(text)
        inferred_points = infer_total_points(text, len(sections))
        st.session_state.detected_assignment_name = inferred_name
        st.session_state.detected_total_points = inferred_points
        if not st.session_state.confirmed_metadata:
            st.session_state.assignment_name = inferred_name
            st.session_state.total_points = inferred_points if inferred_points else 100
        st.success("PDF loaded successfully.")
    except Exception as exc:
        st.error(f"Could not read the PDF: {exc}")
elif manual_text.strip():
    text = clean_text(manual_text)
    st.session_state.source_text = text
    st.session_state.source_loaded = True
    sections = split_into_problems(text)
    inferred_name = infer_assignment_name(text)
    inferred_points = infer_total_points(text, len(sections))
    st.session_state.detected_assignment_name = inferred_name
    st.session_state.detected_total_points = inferred_points
    if not st.session_state.confirmed_metadata:
        st.session_state.assignment_name = inferred_name
        st.session_state.total_points = inferred_points if inferred_points else 100
else:
    st.info("Upload a PDF or paste assignment text to generate a rubric.")

text = st.session_state.source_text

if text:
    st.markdown("### 2 Confirm Assignment Details")
    col1, col2 = st.columns(2)
    with col1:
        assignment_name = st.text_input(
            "Assignment name",
            value=st.session_state.assignment_name,
            placeholder="Enter assignment title",
        )
    with col2:
        detected_points = int(st.session_state.total_points) if st.session_state.total_points else 100
        total_points = st.number_input(
            "Total points",
            min_value=10,
            max_value=500,
            value=detected_points,
            step=5,
        )

    st.caption(
        f"Detected title: {st.session_state.detected_assignment_name or 'Not detected'} | Detected total points: {st.session_state.detected_total_points if st.session_state.detected_total_points else 'Not detected'}"
    )

    confirm_col1, confirm_col2 = st.columns([1, 4])
    with confirm_col1:
        if st.button("Confirm details"):
            st.session_state.assignment_name = assignment_name.strip()
            st.session_state.total_points = int(total_points)
            st.session_state.confirmed_metadata = True
    with confirm_col2:
        if st.session_state.confirmed_metadata:
            st.success("Assignment details confirmed.")

    if show_raw_text:
        with st.expander("Extracted assignment text", expanded=False):
            st.text(text)

    if st.session_state.confirmed_metadata:
        assignment_name = st.session_state.assignment_name or "Untitled Assignment"
        total_points = int(st.session_state.total_points)
        problems = parse_problems(text, total_points)
        default_selection = infer_default_value_set(text)

        st.markdown("### 3 Choose VALUE Rubrics")
        selected_rubrics = st.multiselect(
            "Rubric families",
            options=list(VALUE_RUBRICS.keys()),
            default=default_selection,
            help="Choose the AAC&U VALUE rubric families that best fit the assignment.",
        )

        if not selected_rubrics:
            st.warning("Select at least one VALUE rubric family.")
            st.stop()

        st.markdown("### 4 Weight the VALUE Rubric Families")
        weight_inputs: Dict[str, int] = {}
        cols = st.columns(len(selected_rubrics))
        for i, rubric_name in enumerate(selected_rubrics):
            with cols[i]:
                weight_inputs[rubric_name] = st.number_input(
                    rubric_name,
                    min_value=0,
                    max_value=100,
                    value=DEFAULT_WEIGHTS.get(rubric_name, 20),
                    step=5,
                    key=f"weight_{rubric_name}",
                )

        problem_types = [p.problem_type for p in problems]
        value_df = build_value_rubric(selected_rubrics, total_points, weight_inputs, assignment_name, text, problem_types)
        problem_map_df = build_problem_map(problems, selected_rubrics)
        grid_df = build_grid_view(value_df)
        markdown_summary = generate_markdown_summary(assignment_name, total_points, problems, value_df)

        tab1, tab2, tab3, tab4 = st.tabs(["Overview", "VALUE Rubric", "Grid View", "Problem Mapping"])

        with tab1:
            st.subheader("Detected Problems")
            overview_df = pd.DataFrame(
                [
                    {
                        "Problem": f"Problem {p.number}",
                        "Title": p.title,
                        "Type": p.problem_type,
                        "Suggested Points": p.suggested_points,
                    }
                    for p in problems
                ]
            )
            st.dataframe(overview_df, use_container_width=True)
            st.subheader("Rubric Summary")
            for rubric_name in selected_rubrics:
                with st.expander(rubric_name, expanded=False):
                    rubric_slice = value_df[value_df["VALUE Rubric"] == rubric_name]
                    for _, row in rubric_slice.iterrows():
                        st.markdown(f"**{row['Criterion']} ({row['Points']} pts)**")
                        st.write(f"Focus: {row['Focus']}")
                        st.markdown(f"- **4:** {row['Level 4']}")
                        st.markdown(f"- **3:** {row['Level 3']}")
                        st.markdown(f"- **2:** {row['Level 2']}")
                        st.markdown(f"- **1:** {row['Level 1']}")

        with tab2:
            st.subheader("AAC&U VALUE-Aligned Rubric")
            st.dataframe(value_df, use_container_width=True)

        with tab3:
            st.subheader("Traditional Rubric Grid")
            st.dataframe(grid_df, use_container_width=True)

        with tab4:
            st.subheader("Problem-to-Rubric Mapping")
            st.dataframe(problem_map_df, use_container_width=True)

        st.divider()
        dl1, dl2, dl3, dl4 = st.columns(4)
        with dl1:
            st.download_button(
                "Download VALUE rubric CSV",
                data=to_csv_bytes(value_df),
                file_name="value_rubric.csv",
                mime="text/csv",
            )
        with dl2:
            st.download_button(
                "Download problem map CSV",
                data=to_csv_bytes(problem_map_df),
                file_name="problem_value_map.csv",
                mime="text/csv",
            )
        with dl3:
            st.download_button(
                "Download rubric markdown",
                data=markdown_summary.encode("utf-8"),
                file_name="value_rubric_summary.md",
                mime="text/markdown",
            )
        with dl4:
            st.download_button(
                "Download rubric DOCX",
                data=to_docx_bytes(assignment_name, total_points, problem_map_df, value_df),
                file_name="value_rubric.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

st.divider()
with st.expander("How this version mirrors AAC&U VALUE"):
    st.write(
        "This version follows the AAC&U VALUE model by organizing performance into rubric families and four achievement levels, while adapting the focus language to the uploaded assignment. It also separates upload, metadata confirmation, rubric selection, and export into a clearer workflow."
    )

st.caption("Run with: streamlit run main.py")
